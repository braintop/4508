#!/usr/bin/env python3
"""Generate Hebrew Word lesson guide for Firebase + React lesson."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTPUT = Path(__file__).parent / "שיעור-Firebase-הוספת-Todo.docx"


def set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = pPr.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi"
    )
    pPr.append(bidi)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    return h


def add_para(doc, text, bold=False, bullet=False):
    style = "List Bullet" if bullet else "Normal"
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def build():
    doc = Document()

    # Default font for Hebrew
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading("שיעור Firebase — הוספת Todo ל-Firestore", 0)
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("מדריך מרצה | React + TypeScript + Firebase")
    set_rtl(sub)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Section 1
    add_heading(doc, "1. מטרת השיעור", 1)
    add_para(doc, "בסוף השיעור התלמידים יבינו:", bullet=True)
    add_para(doc, "מה זה Firestore ואיך הוא שונה מ-SQL", bullet=True)
    add_para(doc, "איך מתחברים ל-Firebase מתוך React", bullet=True)
    add_para(doc, "איך מגדירים טיפוסים (TypeScript) למסמכים", bullet=True)
    add_para(doc, "איך מוסיפים רשומה חדשה עם addDoc", bullet=True)
    add_para(doc, "איך לראות את הנתונים ב-Firebase Console", bullet=True)

    # Section 2
    add_heading(doc, "2. הכנה לפני השיעור (מרצה)", 1)
    add_para(doc, "Firebase Console — צור פרויקט חדש (או השתמש בקיים)", bullet=True)
    add_para(doc, "Build → Firestore Database → Create database (Test mode לשיעור)", bullet=True)
    add_para(doc, "Project Settings → Your apps → Web app → העתק את firebaseConfig", bullet=True)
    add_para(doc, "הדבק את הערכים בקובץ .env (בשורש my-first-react-ts)", bullet=True)
    add_para(doc, "ודא ש-.env נמצא ב-.gitignore (לא מעלים סודות ל-Git!)", bullet=True)
    add_para(doc, "הרצה: cd my-first-react-ts && npm install && npm run dev", bullet=True)

    add_para(doc, "משתני סביבה ב-.env:", bold=True)
    add_code(
        doc,
        "VITE_FIREBASE_API_KEY=...\n"
        "VITE_FIREBASE_AUTH_DOMAIN=...\n"
        "VITE_FIREBASE_PROJECT_ID=...\n"
        "VITE_FIREBASE_STORAGE_BUCKET=...\n"
        "VITE_FIREBASE_MESSAGING_ID=...\n"
        "VITE_FIREBASE_APP_ID=...",
    )

    # Section 3 - Architecture
    add_heading(doc, "3. תמונה כללית — זרימת הנתונים", 1)
    add_para(
        doc,
        "הסבירו לתלמידים את השרשרת: "
        "משתמש מקליד טקסט → React שולח אובייקט → Firestore שומר מסמך → "
        "Firebase Console מציג את התוצאה.",
    )
    add_code(
        doc,
        "AddTodo.tsx  →  addDoc()  →  todosCollection  →  Firestore (ענן)\n"
        "                                    ↑\n"
        "                              collections.ts\n"
        "                                    ↑\n"
        "                               firebase.ts (db)",
    )

    # Section 4 - File order
    add_heading(doc, "4. סדר קבצים להסבר (מהתחתית למעלה)", 1)
    add_para(
        doc,
        "הסבירו מהשכבה הנמוכה לגבוהה — כך התלמיד מבין למה כל קובץ קיים לפני שרואה UI.",
    )

    files = [
        (
            "שלב 1 — .env",
            "מפתחות החיבור לפרויקט Firebase. "
            "Vite קורא משתנים שמתחילים ב-VITE_. "
            "לא שומרים את הקובץ ב-Git.",
            None,
        ),
        (
            "שלב 2 — src/types/firebase.ts",
            "מגדירים את צורת המסמך — מה Firestore אמור להכיל. "
            "Todo = text, completed, createdAt, userId + id. "
            "NewTodo = Todo בלי id (כי Firestore יוצר id אוטומטית).",
            "export interface Todo {\n"
            "  id: string\n"
            "  text: string\n"
            "  completed: boolean\n"
            "  createdAt: Date\n"
            "  userId: string\n"
            "}\n\n"
            "export type NewTodo = Omit<Todo, 'id'>",
        ),
        (
            "שלב 3 — src/firebase.ts",
            "חיבור חד-פעמי ל-Firebase. initializeApp + getFirestore. "
            "מייצאים db (וגם auth, storage לשיעורים הבאים). "
            "כל האפליקציה משתמשת באותו db.",
            "export const db = getFirestore(app)",
        ),
        (
            "שלב 4 — src/collections.ts",
            "הפניה מוקלדת (typed) לאוסף todos. "
            "collection(db, 'todos') — שם האוסף בענן. "
            "withConverter — מתרגם בין TypeScript ל-Firestore, "
            "ומוסיף id בקריאה.",
            "export const todosCollection = collection(db, 'todos')\n"
            "  .withConverter(createConverter<Todo>())",
        ),
        (
            "שלב 5 — src/Components/AddTodo.tsx",
            "הממשק + הלוגיקה. useState לטקסט. "
            "בלחיצה: בונים NewTodo, קוראים addDoc, מקבלים id חדש.",
            "const docRef = await addDoc(todosCollection, newTodo)\n"
            "setMessage(`נוסף בהצלחה! id: ${docRef.id}`)",
        ),
        (
            "שלב 6 — src/App.tsx",
            "מרכיב את הקומפוננטה בדף. שורה אחת — <AddTodo />.",
            None,
        ),
    ]

    for title, explanation, code in files:
        add_heading(doc, title, 2)
        add_para(doc, explanation)
        if code:
            add_para(doc, "דוגמת קוד:", bold=True)
            add_code(doc, code)

    # Section 5 - Demo
    add_heading(doc, "5. הדגמה חיה — צעד אחר צעד", 1)
    steps = [
        "פתחו טרמינל בתיקייה my-first-react-ts (לא בתיקיית האב!)",
        "הריצו: npm run dev",
        "פתחו בדפדפן את הכתובת (למשל http://localhost:5173)",
        "הקלידו todo ולחצו 'הוסף'",
        "הראו הודעת הצלחה עם id",
        "פתחו Firebase Console → Firestore → אוסף todos",
        "הראו את המסמך החדש עם כל השדות",
        "הוסיפו עוד todo — הראו שכל מסמך מקבל id ייחודי",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"{i}. {step}", bullet=True)

    # Section 6 - Key concepts
    add_heading(doc, "6. מושגי מפתח להדגשה", 1)
    concepts = [
        ("Collection (אוסף)", "כמו 'טבלה' — todos מכיל מסמכי Todo"),
        ("Document (מסמך)", "רשומה בודדת עם id ייחודי"),
        ("addDoc", "יוצר מסמך חדש — Firestore מייצר id"),
        ("Converter", "גשר בין TypeScript ל-Firestore + הוספת id בקריאה"),
        ("NewTodo vs Todo", "בכתיבה אין id; בקריאה יש id"),
        ("async/await", "פעולות ענן לוקחות זמן — מחכים לתשובה"),
    ]
    for term, meaning in concepts:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        r1 = p.add_run(f"{term}: ")
        r1.bold = True
        r1.font.name = "Arial"
        r2 = p.add_run(meaning)
        r2.font.name = "Arial"

    # Section 7 - Errors
    add_heading(doc, "7. שגיאות נפוצות ופתרונות", 1)
    errors = [
        (
            "ENOENT package.json",
            "הרצתם npm מתיקייה הלא נכונה. עברו ל-my-first-react-ts.",
        ),
        (
            "Missing or insufficient permissions",
            "Firestore Rules חוסמות כתיבה. לשיעור: allow read, write: if true",
        ),
        (
            "Firebase: Error (auth/invalid-api-key)",
            "בדקו ש-.env תקין ושהאפליקציה אותחלה מחדש אחרי שינוי .env",
        ),
        (
            "הטופס לא שולח",
            "בדקו Console בדפדפן (F12). ודאו ש-Firestore Database נוצר",
        ),
    ]
    for err, fix in errors:
        p = doc.add_paragraph()
        set_rtl(p)
        r1 = p.add_run(f"❌ {err}\n")
        r1.bold = True
        r1.font.name = "Arial"
        r2 = p.add_run(f"✅ {fix}")
        r2.font.name = "Arial"

    # Section 8 - Firestore rules
    add_heading(doc, "8. Firestore Rules (לשיעור בלבד)", 1)
    add_para(
        doc,
        "הסבירו שזה רק לפיתוח! בפרודקשן חייבים rules מגבילות.",
    )
    add_code(
        doc,
        "rules_version = '2';\n"
        "service cloud.firestore {\n"
        "  match /databases/{database}/documents {\n"
        "    match /{document=**} {\n"
        "      allow read, write: if true;\n"
        "    }\n"
        "  }\n"
        "}",
    )

    # Section 9 - Homework / next lesson
    add_heading(doc, "9. שיעור הבא (תצוגה מקדימה)", 1)
    add_para(doc, "קריאת todos מהמסד (getDocs / onSnapshot)", bullet=True)
    add_para(doc, "הצגת רשימה ב-UI", bullet=True)
    add_para(doc, "סימון completed + עדכון (updateDoc)", bullet=True)
    add_para(doc, "חיבור auth — userId אמיתי במקום demo-user", bullet=True)

    # Section 10 - Tree
    add_heading(doc, "10. מבנה תיקיות הפרויקט", 1)
    add_code(
        doc,
        "my-first-react-ts/\n"
        "├── .env                    ← מפתחות Firebase\n"
        "├── package.json\n"
        "└── src/\n"
        "    ├── types/firebase.ts   ← טיפוסים\n"
        "    ├── firebase.ts         ← חיבור\n"
        "    ├── collections.ts      ← אוספים\n"
        "    ├── Components/\n"
        "    │   └── AddTodo.tsx     ← טופס הוספה\n"
        "    ├── App.tsx             ← דף ראשי\n"
        "    └── main.tsx            ← נקודת כניסה",
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("קורס 4508 | Firebase + React + TypeScript")
    set_rtl(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
